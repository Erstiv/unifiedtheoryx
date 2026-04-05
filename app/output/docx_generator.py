"""DOCX Export — editable Word documents for papers and scripts."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


NAVY = RGBColor(0x1a, 0x1f, 0x36)
AMBER = RGBColor(0xd4, 0xa5, 0x74)
DARK = RGBColor(0x2c, 0x2c, 0x2c)
MUTED = RGBColor(0x6b, 0x72, 0x80)


def generate_paper_docx(title: str, subtitle: str, content: str, output_path: str):
    """Generate an editable DOCX for the blog post."""
    doc = Document()

    # Style setup
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    font.color.rgb = DARK

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.runs[0].font.color.rgb = NAVY

    if subtitle:
        sub_para = doc.add_paragraph(subtitle)
        sub_para.runs[0].font.italic = True
        sub_para.runs[0].font.color.rgb = MUTED
        sub_para.runs[0].font.size = Pt(13)

    series_para = doc.add_paragraph("The Grand Unified Theory of X")
    series_para.runs[0].font.color.rgb = AMBER
    series_para.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # Content
    if content:
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            # Section headings
            if line.startswith("## "):
                heading = doc.add_heading(line[3:].strip(), level=2)
                heading.runs[0].font.color.rgb = NAVY
                continue

            if line.startswith("# "):
                heading = doc.add_heading(line[2:].strip(), level=1)
                heading.runs[0].font.color.rgb = NAVY
                continue

            # Blockquotes
            if line.startswith("> "):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.right_indent = Inches(0.5)
                run = para.add_run(line[2:].strip())
                run.font.italic = True
                run.font.color.rgb = AMBER
                run.font.size = Pt(12)
                continue

            # Regular paragraph with inline formatting
            para = doc.add_paragraph()
            _add_formatted_text(para, line)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("--- The Grand Unified Theory of X ---")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = MUTED
    footer.runs[0].font.size = Pt(9)

    doc.save(output_path)


def generate_script_docx(title: str, script: str, output_path: str):
    """Generate a screenplay-format DOCX for the podcast script."""
    doc = Document()

    # Courier-style for scripts
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = Pt(10)
    font.color.rgb = DARK

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.name = "Courier New"

    sub = doc.add_paragraph("The Grand Unified Theory of X — Podcast Script")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = MUTED

    doc.add_paragraph()

    if script:
        for line in script.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # Timing markers
            if line.startswith("[TIMING:"):
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.runs[0].font.color.rgb = AMBER
                para.runs[0].font.size = Pt(9)
                continue

            # Directions
            if line.startswith("[DIRECTION:") or line.startswith("[SFX:"):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(line)
                run.font.italic = True
                run.font.color.rgb = MUTED
                run.font.size = Pt(9)
                continue

            # Character names
            char_match = re.match(r"\[([A-Z]+)\]:?\s*(.*)", line)
            if char_match:
                char_name = char_match.group(1)
                dialogue = char_match.group(2)

                name_para = doc.add_paragraph()
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = name_para.add_run(char_name)
                run.font.bold = True
                run.font.color.rgb = NAVY

                if dialogue:
                    dial_para = doc.add_paragraph()
                    dial_para.paragraph_format.left_indent = Inches(1)
                    dial_para.paragraph_format.right_indent = Inches(1)
                    dial_para.add_run(dialogue)
                continue

            # Regular dialogue
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(1)
            para.paragraph_format.right_indent = Inches(1)
            para.add_run(line)

    doc.save(output_path)


def _add_formatted_text(para, text: str):
    """Parse inline markdown bold/italic and add formatted runs."""
    # Split on bold and italic markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.font.italic = True
        else:
            para.add_run(part)
